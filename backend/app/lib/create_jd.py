import openai
import io
import os
import datetime
import asyncio
import uuid
from docx import Document
from docx2python import docx2python as HtmlToDocx
# from builtins import PendingDeprecationWarning
from azure.identity.aio import DefaultAzureCredential
from azure.storage.blob.aio import BlobServiceClient
from azure.cosmos.aio import CosmosClient
from dotenv import load_dotenv
 
load_dotenv()
 
ADLS_CONTAINER = 'landing'
SAS_TOKEN = os.environ["SAS_TOKEN"]
COSMOS_URL = os.environ["COSMOS_URL"]
COSMOS_KEY = os.environ["COSMOS_KEY"]
COSMOS_DATABASE = os.environ["COSMOS_DATABASE"]
STORAGE_ACCOUNT_URL = os.environ["STORAGE_ACCOUNT_URL"]
CONTAINER_NAME = 'recruitbot'
 
async def insert_record(data:dict, container_name) -> None:
    async with CosmosClient(COSMOS_URL, credential=COSMOS_KEY, connection_verify=False) as client:
        database = client.get_database_client(COSMOS_DATABASE)
        container = database.get_container_client(container_name)        
        await container.upsert_item(data)
 
async def create_and_upload_blob_file(html_content, params):
 
    document = Document()
    # document.add_picture('static/logo/Adani Logo.png')
    document.add_paragraph(html_content, style='Normal')  
    # new_parser = HtmlToDocx()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Adani Regular'
    # new_parser.add_html_to_document(html_content, document)
 
    io_buffer = io.BytesIO()
    document.save(io_buffer)
    io_buffer.seek(0)
 
    filename = f"{params['user_name']}/jd/{datetime.date.today().strftime('%Y%m%d')}/{str(params['job_title']).strip()}_{datetime.datetime.now().strftime('%H%M%S')}.docx"
 
    async def upload_blob():  # Define an inner async function
        async with DefaultAzureCredential() as credential:
            async with BlobServiceClient(STORAGE_ACCOUNT_URL, credential, connection_verify=False) as blob_service_client:
                container_client = blob_service_client.get_container_client(container=ADLS_CONTAINER)
                blob_client = await container_client.upload_blob(
                    name=filename, data=io_buffer, overwrite=True
                )
                return blob_client.url
 
    params['file_created'] = await upload_blob()
    params['id'] = str(uuid.uuid4())
    params['run_date'] = str(datetime.datetime.now().date())
   
    await insert_record(params, container_name='recruitbot')
 
    return params['file_created']
 
def create_prompt(parameters: dict) -> str:
    """ Create a prompt for the job description GPT-3 model.
        Args:
        parameters (dict): A dictionary containing the following
        keys: 'job_title', 'industry', 'location', 'experience', 'educational_qualification', 'certifications', 'Technical_skills', 'behavioural_skills', 'key_responsibilities'.
 
        Returns:
        str: A prompt for the job description GPT-3 model.
     """
    if 'key_skills' in parameters.keys() and parameters['key_skills'] != '':
        key_skills = f"with non-negotiable key skills: {parameters['key_skills']}, "
    else:
        key_skills = ""
   
    if 'educational_qualification' in parameters.keys() and parameters['educational_qualification'] != '':
        educational_qualification = f"with Education Qualification as: {parameters['educational_qualification']}, "
    else:
        educational_qualification = ""
   
    if 'key_responsibilities' in parameters.keys() and parameters['key_responsibilities'] != '':
        key_responsibilities = f"with Roles and Responsibilities as: {parameters['key_responsibilities']}, "
    else:
        key_responsibilities = ""
   
    if 'behavioural_skills' in parameters.keys() and parameters['behavioural_skills'] != '':
        behavioural_skills = f"with Behavioural Skills as: {parameters['behavioural_skills']}, "
    else:
        behavioural_skills = ""
   
    if 'certifications' in parameters.keys() and parameters['certifications'] != '':
        certifications = f"with Certifications: {parameters['certifications']}, "
    else:
        certifications = ""
   
    if 'additional_inputs' in parameters.keys() and parameters['additional_inputs'] != '':
        additional_inputs = f"also include: {parameters['additional_inputs']} non-negotiable roles, "
    else:
        additional_inputs = ""
    prompt = f"""
Write comprehensive job description of {parameters['job_title']} having {parameters['experience']} years \
of experience in {parameters['industry']} industry with location in {parameters['location']}, {key_skills}{educational_qualification} \
{key_responsibilities}{behavioural_skills}{certifications}{additional_inputs}. You must elaborate the Roles and \
Responsibilities with at least 500 words. Your generated job description should have summary section. Consider \
the Adani company culture and values, ensuring the job description template is tailored to attract top talent \
while providing a clear understanding of the expectations and opportunities associated with the position.\
Also include a brief conclusion in the end. Convert the complete description in HTML format with all \
headings h4 only and list sub items as bullet points where necessary. You should strictly include all the description in the html otherwise you'll be penalized. \
Do not give bullet points in summary section and do not include years of experience in Non Negotiable skills section"""
    return prompt
 
 
async def get_job_description(user_prompt: str, params):
 
    system_role = """You are an HR who specialises in Talent acquistion for different roles across different industries like port,manufacturing, distribution,retail,mining,energy,infrastructre
            Automobile,design,Construction,Information technology,Transmission,Renewable energy,Oil & Gas, Defence,Arms and Ammunition,Drone system,Missile Design & System, Aerospace, arms industry,chemical, Iron & steel,Banking and Finance,sports,Foodprocessing.You have 20 years of experience in this role which encompasses
            having regular discussuions with the business and understanding the business requirements.You specializes in constructing comphrensive JD (Job Description) which satisfies all the specific inputs given by the user.
            Your job description should includes Job title,Location,Summary,Experience,Roles and Resposibilities,Education Qualification,Certification required,Behavioural skills,Technical Skills,Three Non-Negotiable Skills.
            You are aware of all the relevant Job description template specific to each kind of role and industry that are avilable in the market.Try to replicate these Job Description templates and maintain
            consistency with the business requirements.Consider the Adani company culture and values, ensuring the job description template is tailored to attract top talent while providing a clear understanding of the expectations
            and opportunities associated with the position. Additionally, incorporate inclusive language and showcase the potential for growth within the Adani organization to appeal to a diverse and motivated pool of candidates
            Specifically mention the technical skills required for the given industry and consider the years of experience while deciding the techinal skills for the Job Description."""
 
 
    # Configure Azure OpenAI Service API
    openai.api_type = "azure"
    openai.api_version = "2024-02-01"
    # openai.api_version = "2023-09-15-preview"
    openai.api_base = os.getenv('OPENAI_API_BASE')
    openai.api_key = os.getenv("OPENAI_API_KEY")
 
 
 
    gptmodel = "gpt-35-turbo"  # Default to GPT-3.5 Turbo line 78
    if params.get('model') == 4:
        gptmodel = "gpt-4"
    # gptmodel = 3.5
 
    if params['model'] == 4:
        gptmodel="ailabgpt-4-turbo"
    else:
        gptmodel="ailabgpt35turbo"
   
    temperature = 0.7
    if params['creativity'] == 'high':
        temperature = 0.9
    elif params['creativity'] == 'medium':
        temperature = 0.7
    elif params['creativity'] == 'low':
        temperature = 0.4
 
    max_tokens = 2000
    if params['token'] == 'small':
        max_tokens = 800
    elif params['token'] == 'medium':
        max_tokens = 1200
    if params['token'] == 'large':
        max_tokens = 2500
    if params['token'] == 'x-large':
        max_tokens = 3500
   
    response = await openai.ChatCompletion.acreate(
        engine = gptmodel,
        #model=gptmodel,
        messages=[
            {"role": "system", "content": system_role},
            {"role": "user", "content": user_prompt}
        ],
        temperature=temperature,
        max_tokens=max_tokens,
        top_p=0.95,
        frequency_penalty=0,
        presence_penalty=0,
        stop=None
    )
 
    # response = await openai.ChatCompletion.acreate(engine=gptmodel,
    #                                                 messages=[
    #                                                     {"role": "system", "content": system_role},
    #                                                     {"role": "user", "content": user_prompt}
    #                                                 ],
    #                                                 temperature=temperature,
    #                                                 max_tokens=max_tokens,
    #                                                 top_p=0.95,
    #                                                 frequency_penalty=0,
    #                                                 presence_penalty=0,
    #                                                 stop=None
    #                                                 )
    job_description = response.choices[0].message.content
    job_description = job_description.replace('```', '').replace('html','').replace("<>","").replace("Three Non-Negotiable Skills","Non-Negotiable Skills").replace('< lang="en">','')
    job_description = job_description.replace('Summary','Purpose')
    return job_description
 
 
async def create_job_description(data):
    jd_prompt = create_prompt(data)
    job_description = await get_job_description(jd_prompt, data)
    result = {
        "message": job_description
    }
    return result
 
 
 